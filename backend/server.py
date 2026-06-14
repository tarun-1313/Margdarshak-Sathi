"""CareerPilot AI - FastAPI backend."""

import os
import io
import json
import uuid
import base64
import logging
import asyncio
from pathlib import Path
from datetime import datetime, timezone, timedelta
from typing import Optional, List, Dict, Any

import httpx
import fitz
import docx as docx_lib
import numpy as np
from dotenv import load_dotenv
from fastapi import FastAPI, APIRouter, Depends, HTTPException, Request, Response, UploadFile, File, Cookie, WebSocket, WebSocketDisconnect
from fastapi.responses import StreamingResponse
from pydantic import BaseModel, Field, EmailStr
from starlette.middleware.cors import CORSMiddleware
from motor.motor_asyncio import AsyncIOMotorClient
import bcrypt
from jose import JWTError, jwt

# Import Gemini client
from gemini_client import gemini_generate_text, gemini_generate_json, gemini_stream_text, GEMINI_MODEL

# Third-party
from fastembed import TextEmbedding
from elevenlabs.client import ElevenLabs
from elevenlabs import VoiceSettings
import websockets as ws_client

from reportlab.lib.pagesizes import LETTER
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib.enums import TA_LEFT, TA_CENTER

# Load env
ROOT = Path(__file__).parent
load_dotenv(ROOT / ".env")

# Try Atlas URL first (MONGODB_URL), fallback to local (MONGO_URL)
MONGO_URL = os.environ.get("MONGODB_URL") or os.environ["MONGO_URL"]
DB_NAME = os.environ["DB_NAME"]
GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY")
YT_KEY = os.environ.get("YOUTUBE_API_KEY")
ADZUNA_APP_ID = os.environ.get("ADZUNA_APP_ID")
ADZUNA_APP_KEY = os.environ.get("ADZUNA_APP_KEY")
RAPIDAPI_KEY = os.environ.get("RAPIDAPI_KEY")
GITHUB_TOKEN = os.environ.get("GITHUB_TOKEN")
ELEVENLABS_API_KEY = os.environ.get("ELEVENLABS_API_KEY")
DEEPGRAM_API_KEY = os.environ.get("DEEPGRAM_API_KEY")

# Custom Auth Env Vars
JWT_SECRET = os.environ.get("JWT_SECRET", "dev-secret-key-change-in-prod-12345")
GOOGLE_CLIENT_ID = os.environ.get("GOOGLE_CLIENT_ID")
GOOGLE_CLIENT_SECRET = os.environ.get("GOOGLE_CLIENT_SECRET")
GITHUB_CLIENT_ID = os.environ.get("GITHUB_CLIENT_ID")
GITHUB_CLIENT_SECRET = os.environ.get("GITHUB_CLIENT_SECRET")
# Frontend URL for OAuth redirects
FRONTEND_URL = os.environ.get("FRONTEND_URL", "http://localhost:3000")

# DB
client = AsyncIOMotorClient(MONGO_URL)
db = client[DB_NAME]

# Logging
logging.basicConfig(level=logging.INFO, format="%(asctime)s %(levelname)s %(name)s - %(message)s")
log = logging.getLogger("careerpilot")

# FastAPI
app = FastAPI(title="CareerPilot AI")
api = APIRouter(prefix="/api")


@app.get("/")
async def root():
    return {"service": "CareerPilot AI", "status": "ok", "version": "1.0.0"}

@api.get("/")
async def api_root():
    return {"service": "CareerPilot AI", "status": "ok", "version": "1.0.0"}


# ---------- helpers ----------
def now_utc() -> datetime:
    return datetime.now(timezone.utc)


def jsonable(doc: dict) -> dict:
    doc.pop("_id", None)
    return doc


async def get_current_user(request: Request, session_token: Optional[str] = Cookie(default=None)) -> dict:
    token = session_token
    if not token:
        auth = request.headers.get("Authorization", "")
        if auth.startswith("Bearer "):
            token = auth[7:]
    if not token:
        raise HTTPException(status_code=401, detail="Not authenticated")
    
    try:
        # Verify JWT token first
        payload = jwt.decode(token, JWT_SECRET, algorithms=["HS256"])
        user_id: str = payload.get("sub")
        if user_id is None:
            raise HTTPException(status_code=401, detail="Invalid token")
    except JWTError:
        pass  # Fall back to checking session in DB (for existing sessions)
        
    sess = await db.user_sessions.find_one({"session_token": token}, {"_id": 0})
    if not sess:
        # If not in sessions, try to use user_id from JWT to get user
        try:
            payload = jwt.decode(token, JWT_SECRET, algorithms=["HS256"])
            user_id = payload.get("sub")
            user = await db.users.find_one({"user_id": user_id}, {"_id": 0})
            if not user:
                raise HTTPException(status_code=401, detail="User not found")
            return user
        except JWTError:
            raise HTTPException(status_code=401, detail="Invalid session")
        
    exp = sess["expires_at"]
    if isinstance(exp, str):
        exp = datetime.fromisoformat(exp)
    if exp.tzinfo is None:
        exp = exp.replace(tzinfo=timezone.utc)
    if exp < now_utc():
        raise HTTPException(status_code=401, detail="Session expired")

    user = await db.users.find_one({"user_id": sess["user_id"]}, {"_id": 0})
    if not user:
        raise HTTPException(status_code=401, detail="User not found")
    return user


# ---------- Pydantic Models ----------
class RegisterIn(BaseModel):
    email: EmailStr
    password: str = Field(..., min_length=8)
    name: Optional[str] = None

class LoginIn(BaseModel):
    email: EmailStr
    password: str

class ProfileIn(BaseModel):
    name: Optional[str] = None
    education: Optional[str] = None
    degree: Optional[str] = None
    graduation_year: Optional[int] = None
    skills: List[str] = []
    interests: List[str] = []
    career_goals: Optional[str] = None
    preferred_industry: Optional[str] = None
    preferred_location: Optional[str] = None
    expected_salary: Optional[str] = None
    github_url: Optional[str] = None
    portfolio_url: Optional[str] = None


class ChatIn(BaseModel):
    message: str
    session_id: Optional[str] = None


class InterviewStartIn(BaseModel):
    role: str
    interview_type: str = "technical"


class InterviewAnswerIn(BaseModel):
    interview_id: str
    answer: str


class SalaryPredictIn(BaseModel):
    role: str
    experience_years: int = 0
    location: str = "India"
    skills: List[str] = []


class VoiceInterviewStart(BaseModel):
    role: str
    interview_type: str = "technical"
    difficulty: str = "intermediate"
    personality: str = "friendly"
    voice: str = "technical_male"
    use_resume: bool = True


class VoiceInterviewAnswer(BaseModel):
    interview_id: str
    transcript: str
    confidence: Optional[int] = None
    duration_s: Optional[float] = None
    wps: Optional[float] = None
    filler_words: Optional[int] = None
    code: Optional[str] = None
    language: Optional[str] = None


class EmotionFrame(BaseModel):
    interview_id: str
    turn_index: int
    emotions: Dict[str, float]


class PublishIn(BaseModel):
    slug: str
    bio: Optional[str] = None
    headline: Optional[str] = None
    show_resume: bool = False
    show_portfolio: bool = True


# ---------- Gemini helper ----------
async def _gemini_json(system: str, user_text: str) -> dict:
    try:
        return await gemini_generate_json(prompt=user_text, system_message=system, temperature=0.2)
    except Exception as e:
        log.warning("Gemini JSON parse failed; error=%s", str(e)[:200])
        return {}


# ---------- Embedding helpers ----------
_embedder: Optional[TextEmbedding] = None
def get_embedder() -> TextEmbedding:
    global _embedder
    if _embedder is None:
        _embedder = TextEmbedding(model_name="BAAI/bge-small-en-v1.5")
        log.info("fastembed loaded")
    return _embedder


async def embed_text(text: str) -> List[float]:
    def _embed():
        vecs = list(get_embedder().embed([text]))
        return vecs[0].tolist() if vecs else []
    return await asyncio.to_thread(_embed)


def cosine(a: List[float], b: List[float]) -> float:
    if not a or not b:
        return 0.0
    av, bv = np.array(a, dtype=np.float32), np.array(b, dtype=np.float32)
    denom = float(np.linalg.norm(av) * np.linalg.norm(bv))
    return float(av @ bv / denom) if denom else 0.0


async def memory_search(user_id: str, query: str, k: int = 5) -> List[dict]:
    try:
        qv = await embed_text(query)
        docs = await db.chat_memory.find({"user_id": user_id}, {"_id": 0}).to_list(500)
        scored = [(cosine(qv, d.get("embedding") or []), d) for d in docs]
        scored.sort(key=lambda x: x[0], reverse=True)
        return [d for _, d in scored[:k]]
    except Exception as e:
        log.warning("memory_search failed: %s", e)
        return []


async def memory_store(user_id: str, role: str, content: str) -> None:
    try:
        vec = await embed_text(content)
        await db.chat_memory.insert_one({
            "user_id": user_id, "role": role, "content": content,
            "embedding": vec, "ts": now_utc().isoformat(),
        })
    except Exception as e:
        log.warning("memory_store failed: %s", e)


# ---------- Voice & TTS ----------
el_client = ElevenLabs(api_key=ELEVENLABS_API_KEY) if ELEVENLABS_API_KEY else None

VOICE_PRESETS = {
    "technical_male": {"voice_id": "JBFqnCBsd6RMkjVDRZzb", "label": "George — Technical Interviewer", "gender": "male"},
    "technical_female": {"voice_id": "EXAVITQu4vr4xnSDxMaL", "label": "Sarah — Senior Engineer", "gender": "female"},
    "hr_female": {"voice_id": "XrExE9yKIg1WjnnlVkGX", "label": "Matilda — HR Recruiter", "gender": "female"},
    "hr_male": {"voice_id": "TX3LPaxmHKxFdv7VOQHJ", "label": "Liam — Recruiter", "gender": "male"},
    "faang_strict": {"voice_id": "onwK4e9ZLuTAKqWW03F9", "label": "Daniel — FAANG Lead", "gender": "male"},
    "startup_friendly": {"voice_id": "pFZP5JQG7iQjIQuC4Bku", "label": "Lily — Startup Manager", "gender": "female"},
}

PERSONALITIES = {
    "friendly": "warm and encouraging; ask follow-ups gently",
    "strict_faang": "rigorous FAANG-style interviewer; probe deeply, ask about edge cases and complexity",
    "startup": "casual startup recruiter; mix of culture and technical questions",
    "hr": "professional HR manager; focus on behavioral and soft skills",
    "technical_architect": "senior technical architect; focus on system design, scalability, trade-offs",
}


@api.post("/voice/tts")
async def voice_tts(payload: Dict[str, str], _: dict = Depends(get_current_user)):
    text = payload.get("text", "").strip()
    voice_key = payload.get("voice", "technical_male")
    if not text:
        raise HTTPException(400, "text required")
    if not el_client:
        raise HTTPException(status_code=503, detail={"code": "tts_unavailable", "message": "ElevenLabs not configured."})
    voice_id = VOICE_PRESETS.get(voice_key, VOICE_PRESETS["technical_male"])["voice_id"]

    def _gen() -> bytes:
        stream = el_client.text_to_speech.convert(
            text=text, voice_id=voice_id, model_id="eleven_turbo_v2_5",
            output_format="mp3_44100_128",
            voice_settings=VoiceSettings(stability=0.5, similarity_boost=0.75, style=0.3, use_speaker_boost=True),
        )
        buf = b""
        for chunk in stream:
            if chunk:
                buf += chunk
        return buf

    audio_bytes = await asyncio.to_thread(_gen)
    b64 = base64.b64encode(audio_bytes).decode()
    return {"audio_b64": b64, "mime": "audio/mpeg", "voice_id": voice_id}


# ---------- auth ----------
def create_session_token(user_id: str) -> str:
    """Create a session token (JWT)."""
    expires_delta = timedelta(days=7)
    to_encode = {
        "sub": user_id,
        "exp": now_utc() + expires_delta,
        "iat": now_utc(),
        "type": "session"
    }
    return jwt.encode(to_encode, JWT_SECRET, algorithm="HS256")

def create_user():
    """Helper to create user objects with default fields."""
    user_id = f"user_{uuid.uuid4().hex[:12]}"
    return {
        "user_id": user_id,
        "email": None,
        "password_hash": None,
        "name": None,
        "picture": None,
        "oauth_provider": None,
        "github_username": None,
        "created_at": now_utc().isoformat(),
        "updated_at": now_utc().isoformat(),
        "onboarded": False,
        "skills": [],
        "interests": [],
    }

async def set_auth_cookie(response: Response, session_token: str):
    is_prod = os.environ.get("ENV") == "production"
    response.set_cookie(
        key="session_token",
        value=session_token,
        max_age=7 * 24 * 3600,
        httponly=True,
        secure=is_prod,
        samesite="lax" if not is_prod else "none",
        path="/",
    )

@api.post("/auth/register")
async def auth_register(body: RegisterIn, response: Response):
    """Register a new user with email/password."""
    existing = await db.users.find_one({"email": body.email.lower()}, {"_id": 0})
    if existing:
        raise HTTPException(status_code=400, detail="User with this email already exists")
    
    password_hash = bcrypt.hashpw(body.password.encode("utf-8"), bcrypt.gensalt())
    user = create_user()
    user.update({
        "email": body.email.lower(),
        "password_hash": password_hash.decode("utf-8"),
        "name": body.name or body.email.split("@")[0],
        "oauth_provider": "email"
    })
    
    await db.users.insert_one(dict(user))
    
    session_token = create_session_token(user["user_id"])
    expires = now_utc() + timedelta(days=7)
    await db.user_sessions.insert_one({
        "user_id": user["user_id"],
        "session_token": session_token,
        "expires_at": expires.isoformat(),
        "created_at": now_utc().isoformat(),
    })
    
    await set_auth_cookie(response, session_token)
    return jsonable(user)


@api.post("/auth/login")
async def auth_login(body: LoginIn, response: Response):
    """Login with email/password."""
    user = await db.users.find_one({"email": body.email.lower()}, {"_id": 0})
    if not user:
        raise HTTPException(status_code=401, detail="Invalid email or password")
    
    if not user.get("password_hash"):
        raise HTTPException(status_code=400, detail="This user uses OAuth to login")
    
    if not bcrypt.checkpw(body.password.encode("utf-8"), user["password_hash"].encode("utf-8")):
        raise HTTPException(status_code=401, detail="Invalid email or password")
    
    session_token = create_session_token(user["user_id"])
    expires = now_utc() + timedelta(days=7)
    await db.user_sessions.insert_one({
        "user_id": user["user_id"],
        "session_token": session_token,
        "expires_at": expires.isoformat(),
        "created_at": now_utc().isoformat(),
    })
    
    await set_auth_cookie(response, session_token)
    return jsonable(user)


@api.get("/auth/google/login")
async def auth_google_login():
    """Get Google OAuth redirect URL."""
    if not GOOGLE_CLIENT_ID:
        raise HTTPException(status_code=501, detail="Google OAuth not configured")
    
    redirect_uri = f"{FRONTEND_URL}/auth/google/callback"
    scopes = ["openid", "email", "profile"]
    google_auth_url = "https://accounts.google.com/o/oauth2/v2/auth"
    params = {
        "client_id": GOOGLE_CLIENT_ID,
        "response_type": "code",
        "redirect_uri": redirect_uri,
        "scope": " ".join(scopes),
        "access_type": "online",
        "prompt": "select_account",
    }
    
    from urllib.parse import urlencode
    auth_url = f"{google_auth_url}?{urlencode(params)}"
    return {"auth_url": auth_url}


@api.post("/auth/google/callback")
async def auth_google_callback(code: str, response: Response):
    """Exchange Google OAuth code for user info and login/register."""
    if not GOOGLE_CLIENT_ID or not GOOGLE_CLIENT_SECRET:
        raise HTTPException(status_code=501, detail="Google OAuth not configured")
    
    token_url = "https://oauth2.googleapis.com/token"
    redirect_uri = f"{FRONTEND_URL}/auth/google/callback"
    token_params = {
        "client_id": GOOGLE_CLIENT_ID,
        "client_secret": GOOGLE_CLIENT_SECRET,
        "code": code,
        "grant_type": "authorization_code",
        "redirect_uri": redirect_uri,
    }
    
    async with httpx.AsyncClient(timeout=20) as hx:
        token_res = await hx.post(token_url, data=token_params)
        if token_res.status_code != 200:
            raise HTTPException(status_code=400, detail="Invalid Google OAuth code")
        
        tokens = token_res.json()
        id_token = tokens.get("id_token")
        if not id_token:
            raise HTTPException(status_code=400, detail="No ID token from Google")
        
        # Get user info
        userinfo_res = await hx.get(
            "https://www.googleapis.com/oauth2/v3/userinfo",
            headers={"Authorization": f"Bearer {tokens['access_token']}"}
        )
        if userinfo_res.status_code != 200:
            raise HTTPException(status_code=400, detail="Couldn't get Google user info")
        
        userinfo = userinfo_res.json()
    
    email = userinfo["email"]
    name = userinfo.get("name", email.split("@")[0])
    picture = userinfo.get("picture")
    
    # Check if user exists
    user = await db.users.find_one({"email": email}, {"_id": 0})
    if not user:
        user = create_user()
        user.update({
            "email": email,
            "name": name,
            "picture": picture,
            "oauth_provider": "google",
        })
        await db.users.insert_one(dict(user))
    else:
        # Update user if they exist
        update_data = {"name": name, "updated_at": now_utc().isoformat()}
        if picture:
            update_data["picture"] = picture
        await db.users.update_one({"user_id": user["user_id"]}, {"$set": update_data})
        user.update(update_data)
    
    session_token = create_session_token(user["user_id"])
    expires = now_utc() + timedelta(days=7)
    await db.user_sessions.insert_one({
        "user_id": user["user_id"],
        "session_token": session_token,
        "expires_at": expires.isoformat(),
        "created_at": now_utc().isoformat(),
    })
    
    await set_auth_cookie(response, session_token)
    return jsonable(user)


@api.get("/auth/github/login")
async def auth_github_login():
    """Get GitHub OAuth redirect URL."""
    if not GITHUB_CLIENT_ID:
        raise HTTPException(status_code=501, detail="GitHub OAuth not configured")
    
    redirect_uri = f"{FRONTEND_URL}/auth/github/callback"
    scopes = ["read:user", "user:email"]
    github_auth_url = "https://github.com/login/oauth/authorize"
    from urllib.parse import urlencode
    params = {
        "client_id": GITHUB_CLIENT_ID,
        "redirect_uri": redirect_uri,
        "scope": " ".join(scopes),
        "allow_signup": "true"
    }
    auth_url = f"{github_auth_url}?{urlencode(params)}"
    return {"auth_url": auth_url}


@api.post("/auth/github/callback")
async def auth_github_callback(code: str, response: Response):
    """Exchange GitHub OAuth code for user info and login/register."""
    if not GITHUB_CLIENT_ID or not GITHUB_CLIENT_SECRET:
        raise HTTPException(status_code=501, detail="GitHub OAuth not configured")
    
    token_url = "https://github.com/login/oauth/access_token"
    redirect_uri = f"{FRONTEND_URL}/auth/github/callback"
    token_params = {
        "client_id": GITHUB_CLIENT_ID,
        "client_secret": GITHUB_CLIENT_SECRET,
        "code": code,
        "redirect_uri": redirect_uri,
    }
    headers = {"Accept": "application/json"}
    
    async with httpx.AsyncClient(timeout=20) as hx:
        token_res = await hx.post(token_url, data=token_params, headers=headers)
        if token_res.status_code != 200:
            raise HTTPException(status_code=400, detail="Invalid GitHub OAuth code")
        
        tokens = token_res.json()
        access_token = tokens.get("access_token")
        if not access_token:
            raise HTTPException(status_code=400, detail="No access token from GitHub")
        
        # Get user info
        userinfo_res = await hx.get(
            "https://api.github.com/user",
            headers={"Authorization": f"token {access_token}"}
        )
        if userinfo_res.status_code != 200:
            raise HTTPException(status_code=400, detail="Couldn't get GitHub user info")
        
        userinfo = userinfo_res.json()
        
        # Get primary email (GitHub might require this)
        email_res = await hx.get(
            "https://api.github.com/user/emails",
            headers={"Authorization": f"token {access_token}"}
        )
        emails = email_res.json() if email_res.status_code == 200 else []
        email = None
        for e in emails:
            if e.get("primary") and e.get("verified"):
                email = e["email"]
                break
        
        if not email and userinfo.get("email"):
            email = userinfo["email"]
        
        if not email:
            raise HTTPException(status_code=400, detail="Couldn't get verified email from GitHub")
    
    name = userinfo.get("name", userinfo.get("login", email.split("@")[0]))
    picture = userinfo.get("avatar_url")
    github_username = userinfo.get("login")
    
    # Check if user exists
    user = await db.users.find_one({"email": email}, {"_id": 0})
    if not user:
        user = create_user()
        user.update({
            "email": email,
            "name": name,
            "picture": picture,
            "oauth_provider": "github",
            "github_username": github_username,
        })
        await db.users.insert_one(dict(user))
    else:
        # Update user
        update_data = {
            "name": name, 
            "github_username": github_username, 
            "updated_at": now_utc().isoformat()
        }
        if picture:
            update_data["picture"] = picture
        await db.users.update_one({"user_id": user["user_id"]}, {"$set": update_data})
        user.update(update_data)
    
    session_token = create_session_token(user["user_id"])
    expires = now_utc() + timedelta(days=7)
    await db.user_sessions.insert_one({
        "user_id": user["user_id"],
        "session_token": session_token,
        "expires_at": expires.isoformat(),
        "created_at": now_utc().isoformat(),
    })
    
    await set_auth_cookie(response, session_token)
    return jsonable(user)


@api.get("/auth/me")
async def auth_me(user: dict = Depends(get_current_user)):
    return user


@api.post("/auth/logout")
async def auth_logout(response: Response, request: Request, session_token: Optional[str] = Cookie(default=None)):
    if session_token:
        await db.user_sessions.delete_one({"session_token": session_token})
    response.delete_cookie("session_token", path="/")
    return {"ok": True}


# ---------- profile ----------
@api.put("/profile")
async def update_profile(body: ProfileIn, user: dict = Depends(get_current_user)):
    updates = {k: v for k, v in body.model_dump().items() if v is not None}
    updates["onboarded"] = True
    updates["updated_at"] = now_utc().isoformat()
    await db.users.update_one({"user_id": user["user_id"]}, {"$set": updates})
    new = await db.users.find_one({"user_id": user["user_id"]}, {"_id": 0})
    return new


# ---------- resume ----------
def _extract_pdf(buf: bytes) -> str:
    text = []
    with fitz.open(stream=buf, filetype="pdf") as doc:
        for page in doc:
            text.append(page.get_text())
    return "\n".join(text)


def _extract_docx(buf: bytes) -> str:
    f = io.BytesIO(buf)
    d = docx_lib.Document(f)
    return "\n".join(p.text for p in d.paragraphs)


@api.post("/resume/upload")
async def resume_upload(file: UploadFile = File(...), user: dict = Depends(get_current_user)):
    buf = await file.read()
    name = (file.filename or "").lower()
    if name.endswith(".pdf"):
        text = _extract_pdf(buf)
    elif name.endswith(".docx"):
        text = _extract_docx(buf)
    else:
        raise HTTPException(400, "Only .pdf or .docx supported")

    if not text.strip():
        raise HTTPException(400, "Could not extract text from resume")

    sys_prompt = (
        "You are an expert resume parser and ATS scorer. Return STRICT JSON only with keys: "
        "skills (list of strings), projects (list of {name, description}), certifications (list of strings), "
        "experience (list of {role, company, duration}), education (list of {degree, institution, year}), "
        "ats_score (int 0-100), strengths (list), improvements (list), missing_keywords (list of strings)."
    )
    parsed = await _gemini_json(sys_prompt, f"Parse and score this resume:\n\n{text[:12000]}")

    resume_id = f"res_{uuid.uuid4().hex[:12]}"
    doc = {
        "resume_id": resume_id,
        "user_id": user["user_id"],
        "filename": file.filename,
        "raw_text": text[:20000],
        "parsed": parsed,
        "ats_score": int(parsed.get("ats_score") or 0),
        "created_at": now_utc().isoformat(),
    }
    await db.resumes.insert_one(dict(doc))

    # update user skills
    if parsed.get("skills"):
        await db.users.update_one(
            {"user_id": user["user_id"]},
            {"$addToSet": {"skills": {"$each": parsed["skills"]}}},
        )
    return jsonable(doc)


@api.get("/resume/latest")
async def resume_latest(user: dict = Depends(get_current_user)):
    doc = await db.resumes.find_one(
        {"user_id": user["user_id"]}, {"_id": 0}, sort=[("created_at", -1)]
    )
    return doc or {}


# ---------- careers ----------
@api.post("/careers/recommend")
async def careers_recommend(user: dict = Depends(get_current_user)):
    sys_prompt = (
        "You are an AI career advisor. Recommend 5 career paths for the user. "
        "Return STRICT JSON: {careers: [{name, match_score (0-100 int), description, "
        "salary_range_inr (string e.g. '12-25 LPA'), demand_level ('Very High'|'High'|'Moderate'), "
        "growth_potential ('Excellent'|'Strong'|'Stable'), key_skills: [string]}]}"
    )
    profile = {
        "skills": user.get("skills", []),
        "interests": user.get("interests", []),
        "education": user.get("education"),
        "degree": user.get("degree"),
        "career_goals": user.get("career_goals"),
        "preferred_industry": user.get("preferred_industry"),
    }
    result = await _gemini_json(sys_prompt, f"User profile:\n{json.dumps(profile)}")
    careers = result.get("careers", [])
    await db.career_recommendations.delete_many({"user_id": user["user_id"]})
    if careers:
        await db.career_recommendations.insert_many([
            {"user_id": user["user_id"], "created_at": now_utc().isoformat(), **c} for c in careers
        ])
    return {"careers": careers}


@api.get("/careers")
async def careers_list(user: dict = Depends(get_current_user)):
    items = await db.career_recommendations.find(
        {"user_id": user["user_id"]}, {"_id": 0}
    ).to_list(20)
    return {"careers": items}


# ---------- skill gap ----------
@api.post("/skills/gap")
async def skill_gap(payload: Dict[str, str], user: dict = Depends(get_current_user)):
    target_role = payload.get("role") or user.get("career_goals") or "AI Engineer"
    sys_prompt = (
        "You are a skill-gap analyst. Return STRICT JSON: "
        "{target_role, current_skills:[string], required_skills:[string], "
        "missing_skills:[{name, priority('high'|'medium'|'low'), difficulty('beginner'|'intermediate'|'advanced'), estimated_weeks:int, reason}]}"
    )
    inp = {"target_role": target_role, "current_skills": user.get("skills", [])}
    result = await _gemini_json(sys_prompt, json.dumps(inp))
    await db.skill_gaps.update_one(
        {"user_id": user["user_id"], "target_role": target_role},
        {"$set": {**result, "user_id": user["user_id"], "target_role": target_role, "created_at": now_utc().isoformat()}},
        upsert=True,
    )
    return result


# ---------- roadmap ----------
@api.post("/roadmap/generate")
async def roadmap_generate(payload: Dict[str, str], user: dict = Depends(get_current_user)):
    target_role = payload.get("role") or user.get("career_goals") or "AI Engineer"
    months = int(payload.get("months") or 6)
    sys_prompt = (
        f"You are a career mentor. Generate a {months}-month personalized learning roadmap. "
        "Return STRICT JSON: {role, months:[{month:int, title:string, focus:string, "
        "weekly_goals:[string], projects:[string], certifications:[string]}]}"
    )
    inp = {"target_role": target_role, "current_skills": user.get("skills", []), "months": months}
    result = await _gemini_json(sys_prompt, json.dumps(inp))
    rid = f"road_{uuid.uuid4().hex[:10]}"
    doc = {
        "roadmap_id": rid, "user_id": user["user_id"], "target_role": target_role,
        "data": result, "completed_milestones": [],
        "created_at": now_utc().isoformat(),
    }
    await db.roadmaps.delete_many({"user_id": user["user_id"]})
    await db.roadmaps.insert_one(dict(doc))
    return jsonable(doc)


@api.get("/roadmap")
async def roadmap_get(user: dict = Depends(get_current_user)):
    doc = await db.roadmaps.find_one({"user_id": user["user_id"]}, {"_id": 0}, sort=[("created_at", -1)])
    return doc or {}


@api.post("/roadmap/milestone/toggle")
async def roadmap_toggle(payload: Dict[str, Any], user: dict = Depends(get_current_user)):
    key = payload.get("key")  # e.g. "1-week-1"
    if not key:
        raise HTTPException(400, "key required")
    doc = await db.roadmaps.find_one({"user_id": user["user_id"]})
    if not doc:
        raise HTTPException(404, "no roadmap")
    completed = set(doc.get("completed_milestones", []))
    if key in completed:
        completed.remove(key)
    else:
        completed.add(key)
    await db.roadmaps.update_one({"user_id": user["user_id"]}, {"$set": {"completed_milestones": list(completed)}})
    return {"completed_milestones": list(completed)}


# ---------- courses ----------
@api.get("/courses/search")
async def courses_search(q: str = "machine learning", user: dict = Depends(get_current_user)):
    items: List[Dict[str, Any]] = []
    if YT_KEY:
        url = "https://www.googleapis.com/youtube/v3/search"
        params = {
            "part": "snippet", "q": f"{q} tutorial course", "type": "video",
            "maxResults": 12, "relevanceLanguage": "en", "key": YT_KEY,
        }
        async with httpx.AsyncClient(timeout=15) as hx:
            r = await hx.get(url, params=params)
        if r.status_code == 200:
            for v in r.json().get("items", []):
                vid = v["id"]["videoId"]
                sn = v["snippet"]
                items.append({
                    "id": vid,
                    "provider": "YouTube",
                    "title": sn["title"],
                    "channel": sn["channelTitle"],
                    "thumbnail": sn["thumbnails"]["high"]["url"],
                    "url": f"https://www.youtube.com/watch?v={vid}",
                    "description": sn.get("description", "")[:200],
                })
    return {"query": q, "results": items}


# ---------- jobs ----------
@api.get("/jobs/search")
async def jobs_search(q: str = "software engineer", location: str = "India", user: dict = Depends(get_current_user)):
    items: List[Dict[str, Any]] = []
    
    # First try JSearch (RapidAPI)
    if RAPIDAPI_KEY:
        try:
            async with httpx.AsyncClient(timeout=20) as hx:
                r = await hx.get(
                    "https://jsearch.p.rapidapi.com/search",
                    headers={"X-RapidAPI-Key": RAPIDAPI_KEY, "X-RapidAPI-Host": "jsearch.p.rapidapi.com"},
                    params={"query": f"{q} in {location}", "page": "1", "num_pages": "1"},
                )
            if r.status_code == 200:
                for j in (r.json().get("data") or [])[:12]:
                    items.append({
                        "id": j.get("job_id"),
                        "title": j.get("job_title"),
                        "company": j.get("employer_name"),
                        "location": ", ".join(filter(None, [j.get("job_city"), j.get("job_country")])),
                        "salary": f"{j.get('job_min_salary')}-{j.get('job_max_salary')} {j.get('job_salary_currency') or ''}" if j.get("job_min_salary") else "Not disclosed",
                        "url": j.get("job_apply_link"),
                        "description": (j.get("job_description") or "")[:280],
                        "remote": j.get("job_is_remote", False),
                        "posted": j.get("job_posted_at_datetime_utc"),
                    })
        except Exception as e:
            log.warning("JSearch failed: %s", e)
    
    # If no results from JSearch, try Adzuna as fallback
    if not items and ADZUNA_APP_ID and ADZUNA_APP_KEY:
        try:
            # Map location to Adzuna country code (simplified)
            country_code = "in"  # Default to India
            loc_lower = location.lower()
            if loc_lower in ["us", "united states", "usa"]: country_code = "us"
            if loc_lower in ["uk", "united kingdom", "england", "scotland", "wales"]: country_code = "gb"
            if loc_lower in ["ca", "canada"]: country_code = "ca"
            if loc_lower in ["au", "australia"]: country_code = "au"
            
            async with httpx.AsyncClient(timeout=20) as hx:
                adzuna_url = f"https://api.adzuna.com/v1/api/jobs/{country_code}/search/1"
                params = {
                    "app_id": ADZUNA_APP_ID,
                    "app_key": ADZUNA_APP_KEY,
                    "results_per_page": 12,
                    "what": q,
                    "where": location,
                }
                r = await hx.get(adzuna_url, params=params)
                if r.status_code == 200:
                    adzuna_data = r.json()
                    for j in (adzuna_data.get("results", []))[:12]:
                        items.append({
                            "id": j.get("id"),
                            "title": j.get("title"),
                            "company": j.get("company", {}).get("display_name", "Unknown Company"),
                            "location": j.get("location", {}).get("display_name", location),
                            "salary": f"{j.get('salary_min', '')}-{j.get('salary_max', '')} {j.get('salary_currency', '')}" if j.get("salary_min") else "Not disclosed",
                            "url": j.get("redirect_url"),
                            "description": (j.get("description", "") or "")[:280],
                            "remote": "remote" in (j.get("title", "") + " " + j.get("description", "")).lower(),
                            "posted": j.get("created"),
                        })
        except Exception as e:
            log.warning("Adzuna failed: %s", e)
            
    # If both fail, generate AI mock jobs as final fallback using Gemini
    if not items:
        try:
            sys_prompt = (
                "You are a job search API. Generate 8 realistic job listings for the given role and location. "
                "Return STRICT JSON with a 'jobs' list where each job has: "
                "title, company, location, salary (string), url (can be fake but valid), description, remote (boolean). "
                "Make the jobs look real, with actual company names and plausible descriptions."
            )
            result = await _gemini_json(sys_prompt, f"Role: {q}, Location: {location}")
            for j in (result.get("jobs", []))[:8]:
                items.append({
                    "id": str(uuid.uuid4()),
                    "title": j.get("title"),
                    "company": j.get("company"),
                    "location": j.get("location"),
                    "salary": j.get("salary", "Not disclosed"),
                    "url": j.get("url", "#"),
                    "description": (j.get("description", "") or "")[:280],
                    "remote": j.get("remote", False),
                    "posted": now_utc().isoformat(),
                })
        except Exception as e:
            log.warning("Gemini mock jobs failed: %s", e)
            
    # User skills matching
    user_skills = {s.lower() for s in user.get("skills", [])}
    for it in items:
        text = (it.get("title", "") + " " + it.get("description", "")).lower()
        hits = sum(1 for s in user_skills if s and s in text)
        it["match_score"] = min(100, 40 + hits * 12) if user_skills else 60

    items.sort(key=lambda x: x["match_score"], reverse=True)
    return {"query": q, "location": location, "results": items}


# ---------- chatbot ----------
@api.post("/chat/stream")
async def chat_stream(body: ChatIn, user: dict = Depends(get_current_user)):
    sid = body.session_id or f"chat_{user['user_id']}"
    profile_ctx = {
        "name": user.get("name"),
        "skills": user.get("skills", []),
        "career_goals": user.get("career_goals"),
        "education": user.get("education"),
    }

    # Persistent semantic recall from MongoDB Atlas
    recall_lines = []
    try:
        hits = await memory_search(user["user_id"], body.message, k=5)
        for h in hits:
            if h.get("role") and h.get("content"):
                recall_lines.append(f"[{h['role']}] {h['content'][:240]}")
    except Exception as e:
        log.warning("memory search err: %s", e)

    memory_block = ("\n".join(recall_lines[:5])) or "(no relevant past memory)"

    sys = (
        "You are CareerPilot AI, a friendly expert career mentor. "
        "Give concise, actionable, modern advice. Use bullet points where useful.\n"
        f"User context: {json.dumps(profile_ctx)}\n"
        f"Relevant past conversation memory:\n{memory_block}"
    )

    await db.chat_messages.insert_one({
        "user_id": user["user_id"], "session_id": sid,
        "role": "user", "content": body.message, "ts": now_utc().isoformat(),
    })
    await memory_store(user["user_id"], "user", body.message)

    async def gen():
        full = []
        try:
            async for chunk in gemini_stream_text(
                prompt=body.message,
                system_message=sys,
                temperature=0.7,
            ):
                full.append(chunk)
                yield f"data: {json.dumps({'delta': chunk})}\n\n"
        except Exception as e:
            yield f"data: {json.dumps({'error': str(e)})}\n\n"
        assistant_text = "".join(full)
        await db.chat_messages.insert_one({
            "user_id": user["user_id"], "session_id": sid,
            "role": "assistant", "content": assistant_text, "ts": now_utc().isoformat(),
        })
        if assistant_text:
            await memory_store(user["user_id"], "assistant", assistant_text)
        yield "data: [DONE]\n\n"

    return StreamingResponse(
        gen(), media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


@api.get("/chat/history")
async def chat_history(user: dict = Depends(get_current_user)):
    sid = f"chat_{user['user_id']}"
    msgs = await db.chat_messages.find(
        {"user_id": user["user_id"], "session_id": sid}, {"_id": 0}
    ).sort("ts", 1).to_list(200)
    return {"messages": msgs}


# ---------- portfolio / github ----------
@api.post("/portfolio/analyze")
async def portfolio_analyze(payload: Dict[str, str], user: dict = Depends(get_current_user)):
    gh_url = payload.get("github_url") or user.get("github_url") or ""
    if "github.com/" not in gh_url:
        raise HTTPException(400, "Provide a valid GitHub URL")
    username = gh_url.rstrip("/").split("github.com/")[-1].split("/")[0]

    headers = {"Accept": "application/vnd.github+json"}
    if GITHUB_TOKEN:
        headers["Authorization"] = f"Bearer {GITHUB_TOKEN}"

    async with httpx.AsyncClient(timeout=20) as hx:
        user_r = await hx.get(f"https://api.github.com/users/{username}", headers=headers)
        repos_r = await hx.get(f"https://api.github.com/users/{username}/repos", headers=headers, params={"sort": "updated", "per_page": 20})
    if user_r.status_code != 200:
        raise HTTPException(404, f"GitHub user '{username}' not found")
    gh_user = user_r.json()
    repos = repos_r.json() if repos_r.status_code == 200 else []

    repo_summary = [{
        "name": r.get("name"),
        "description": r.get("description"),
        "language": r.get("language"),
        "stars": r.get("stargazers_count", 0),
        "forks": r.get("forks_count", 0),
        "url": r.get("html_url"),
        "has_pages": r.get("has_pages", False),
    } for r in repos[:15]]

    sys_prompt = (
        "You are a senior engineer reviewing a GitHub portfolio. Return STRICT JSON: "
        "{score:int 0-100, summary:string, strengths:[string], improvements:[string], top_projects:[{name, comment}]}"
    )
    profile_text = {
        "user": {"login": gh_user.get("login"), "name": gh_user.get("name"), "bio": gh_user.get("bio"), "followers": gh_user.get("followers"), "public_repos": gh_user.get("public_repos")},
        "repos": repo_summary,
    }
    ai = await _gemini_json(sys_prompt, json.dumps(profile_text))

    out = {
        "username": username, "avatar": gh_user.get("avatar_url"), "name": gh_user.get("name") or username,
        "bio": gh_user.get("bio"), "public_repos": gh_user.get("public_repos", 0), "followers": gh_user.get("followers", 0),
        "repos": repo_summary, "analysis": ai, "created_at": now_utc().isoformat(),
    }
    await db.portfolios.update_one({"user_id": user["user_id"]}, {"$set": {"user_id": user["user_id"], **out}}, upsert=True)
    if gh_url:
        await db.users.update_one({"user_id": user["user_id"]}, {"$set": {"github_url": gh_url}})
    return out


@api.get("/portfolio")
async def get_portfolio(user: dict = Depends(get_current_user)):
    portfolio = await db.portfolios.find_one({"user_id": user["user_id"]}, {"_id": 0})
    if portfolio:
        if "avatar" in portfolio and "avatar_url" not in portfolio:
            portfolio["avatar_url"] = portfolio["avatar"]
        if "repos" in portfolio:
            for repo in portfolio["repos"]:
                if "url" in repo and "html_url" not in repo:
                    repo["html_url"] = repo["url"]
    return portfolio if portfolio else None


@api.get("/resume/latest")
async def get_latest_resume(user: dict = Depends(get_current_user)):
    latest = await db.resumes.find_one({"user_id": user["user_id"]}, {"_id": 0}, sort=[("created_at", -1)])
    return latest if latest else None



# ---------- salary ----------
@api.post("/salary/predict")
async def salary_predict(payload: SalaryPredictIn, user: dict = Depends(get_current_user)):
    sys_prompt = (
        "You are a compensation analyst. Return STRICT JSON: "
        "{role, experience_years, india_range_inr:string, remote_range_usd:string, "
        "international_range_usd:string, factors:[string], confidence(0-100 int)}"
    )
    inp = payload.model_dump()
    inp["user_skills"] = user.get("skills", [])
    res = await _gemini_json(sys_prompt, json.dumps(inp))
    return res


# ---------- trends ----------
@api.get("/trends")
async def trends(user: dict = Depends(get_current_user)):
    sys_prompt = (
        "Return STRICT JSON of industry trends for 2026: "
        "{trending_tech:[{name, momentum:int 1-100, category}], "
        "high_demand_roles:[{role, demand:int, growth:string}], "
        "salary_shifts:[{role, change_pct:int, note}], "
        "ai_focus_areas:[string]}"
    )
    res = await _gemini_json(sys_prompt, "Generate current tech industry trends.")
    return res

# ---------- progress ----------
@api.get("/progress")
async def progress(user: dict = Depends(get_current_user)):
    # Roadmap progress
    roadmap = await db.roadmaps.find_one({"user_id": user["user_id"]}, {"_id": 0})
    roadmap_total = 0
    roadmap_done = 0
    if roadmap and roadmap.get("data", {}).get("months"):
        months = roadmap["data"]["months"]
        roadmap_total = len(months) * 4  # approx 4 weeks per month
        roadmap_done = len(roadmap.get("completed_milestones", []))
    roadmap_percent = int((roadmap_done / roadmap_total * 100) if roadmap_total > 0 else 0)
    
    # Interviews taken
    interviews_taken = await db.interviews.count_documents({"user_id": user["user_id"]})
    
    # Resumes uploaded
    resumes_uploaded = await db.resumes.count_documents({"user_id": user["user_id"]})
    
    # ATS score
    latest_resume = await db.resumes.find_one({"user_id": user["user_id"]}, {"_id": 0}, sort=[("created_at", -1)])
    ats_score = latest_resume.get("ats_score", 0) if latest_resume else 0
    
    # Skills count
    skills_count = len(user.get("skills", []))
    
    return {
        "roadmap_total": roadmap_total,
        "roadmap_done": roadmap_done,
        "roadmap_percent": roadmap_percent,
        "interviews_taken": interviews_taken,
        "resumes_uploaded": resumes_uploaded,
        "ats_score": ats_score,
        "skills_count": skills_count
    }


# ---------- analytics (Complete AI Intelligence Dashboard) ----------
@api.get("/analytics/dashboard")
async def get_dashboard_analytics(user: dict = Depends(get_current_user)):
    # --- 0. Initialize or Fetch User Analytics Record ---
    analytics = await db.analytics.find_one({"user_id": user["user_id"]})
    if not analytics:
        analytics = {
            "user_id": user["user_id"],
            "total_profile_views": 0,
            "recruiter_profile_views": 0,
            "ai_activity_feed": [],
            "created_at": now_utc().isoformat(),
            "updated_at": now_utc().isoformat()
        }
        await db.analytics.insert_one(analytics)

    # --- 1. ATS Score from Resume ---
    latest_resume = await db.resumes.find_one({"user_id": user["user_id"]}, {"_id": 0}, sort=[("created_at", -1)])
    ats_score = latest_resume.get("ats_score", 0) if latest_resume else 0

    # --- 2. Interviews Count & Readiness ---
    interviews_taken = await db.interviews.count_documents({"user_id": user["user_id"]})
    interview_readiness = min(60 + (interviews_taken * 8), 95) if interviews_taken > 0 else 55

    # --- 3. Profile Views & Traffic ---
    total_views = analytics.get("total_profile_views", 0)
    recruiter_views = analytics.get("recruiter_profile_views", 0)
    profile_views_growth = analytics.get("views_growth", 0)

    # --- 4. Roadmap Progress ---
    roadmap = await db.roadmaps.find_one({"user_id": user["user_id"]}, {"_id": 0})
    roadmap_total = 0
    roadmap_done = 0
    if roadmap and roadmap.get("data", {}).get("months"):
        months = roadmap["data"]["months"]
        roadmap_total = len(months) * 4
        roadmap_done = len(roadmap.get("completed_milestones", []))
    roadmap_percent = int((roadmap_done / roadmap_total * 100) if roadmap_total > 0 else 0)

    # --- 5. GitHub & Portfolio Analysis ---
    portfolio = await db.portfolios.find_one({"user_id": user["user_id"]}, {"_id": 0})
    github_score = portfolio.get("analysis", {}).get("score", 0) if portfolio else 0
    github_repos = portfolio.get("public_repos", 0) if portfolio else 0
    github_stars = sum([r.get("stars", 0) for r in portfolio.get("repos", [])]) if portfolio else 0

    # --- 6. AI Skill Score ---
    user_skills = user.get("skills", [])
    ai_skill_score = min(30 + len(user_skills) * 10 + (github_score // 3), 100)

    # --- 7. AI Activity Feed ---
    ai_feed = []
    if latest_resume and ats_score > 0:
        ai_feed.append({
            "icon": "resume",
            "title": "Resume ATS Score Updated",
            "description": f"Your current ATS score is {ats_score}/100",
            "time": "Just now",
            "type": "success"
        })

    if portfolio:
        ai_feed.append({
            "icon": "github",
            "title": "GitHub Portfolio Analyzed",
            "description": f"Portfolio quality score: {github_score}/100",
            "time": "2 hours ago",
            "type": "info"
        })

    if interviews_taken > 0:
        ai_feed.append({
            "icon": "interview",
            "title": "Mock Interview Completed",
            "description": f"You've completed {interviews_taken} interview sessions",
            "time": "3 hours ago",
            "type": "success"
        })

    # --- 8. Skill Radar Data ---
    skill_radar = [
        {"name": "Python", "value": 30 + len(user_skills)},
        {"name": "AI/ML", "value": 25 + (github_score // 4)},
        {"name": "Backend", "value": 40},
        {"name": "Frontend", "value": 35},
        {"name": "DSA", "value": 30 + (interviews_taken * 5)},
        {"name": "Cloud", "value": 20}
    ]

    # --- 9. Job Matches ---
    job_matches = [
        {"role": "AI Engineer", "match": 85 + (ai_skill_score // 10), "company": "Tech Corp", "salary": "25-40 LPA"},
        {"role": "ML Engineer", "match": 80 + (ai_skill_score // 10), "company": "Data Co", "salary": "20-35 LPA"}
    ]

    return {
        "ats_score": ats_score,
        "interview_readiness": interview_readiness,
        "interviews_taken": interviews_taken,
        "profile_views": total_views,
        "recruiter_views": recruiter_views,
        "profile_views_growth": profile_views_growth,
        "roadmap_percent": roadmap_percent,
        "roadmap_done": roadmap_done,
        "roadmap_total": roadmap_total,
        "ai_skill_score": ai_skill_score,
        "github_score": github_score,
        "github_repos": github_repos,
        "github_stars": github_stars,
        "skill_radar": skill_radar,
        "ai_feed": ai_feed if ai_feed else [{"icon": "sparkles", "title": "Welcome to CareerPilot!", "description": "Upload your resume to get started", "time": "Just now", "type": "info"}],
        "job_matches": job_matches,
        "user_skills": user_skills,
        "last_updated": now_utc().isoformat()
    }

# ---------- interview ----------
class InterviewStartIn(BaseModel):
    role: str
    interview_type: str = "technical"

class InterviewAnswerIn(BaseModel):
    interview_id: str
    answer: str

@api.post("/interview/start")
async def interview_start(payload: InterviewStartIn, user: dict = Depends(get_current_user)):
    interview_id = f"int_{uuid.uuid4().hex[:12]}"
    sys_prompt = f"You are a {payload.interview_type} interviewer for the role of {payload.role}. Ask the first question. Return STRICT JSON with 'question': string"
    res = await _gemini_json(sys_prompt, "Ask the first interview question.")
    question = res.get("question", "Tell me about yourself.")
    doc = {
        "interview_id": interview_id,
        "user_id": user["user_id"],
        "role": payload.role,
        "interview_type": payload.interview_type,
        "status": "active",
        "created_at": now_utc().isoformat(),
        "qa": [{"q": question}]
    }
    await db.interviews.insert_one(dict(doc))
    return jsonable(doc)

@api.post("/interview/answer")
async def interview_answer(payload: InterviewAnswerIn, user: dict = Depends(get_current_user)):
    interview = await db.interviews.find_one({"interview_id": payload.interview_id, "user_id": user["user_id"]})
    if not interview:
        raise HTTPException(status_code=404, detail="Interview not found")
    # Update last QA with answer
    qa = interview["qa"]
    qa[-1]["a"] = payload.answer
    
    # Ask next question or finish
    sys_prompt = f"You are a {interview['interview_type']} interviewer for {interview['role']}. Given the conversation history, ask the next question. If you think we've asked enough questions, return {{'finished': true, 'feedback': string}}. Otherwise return {{'question': string}}."
    history = json.dumps(qa)
    res = await _gemini_json(sys_prompt, history)
    if res.get("finished"):
        await db.interviews.update_one({"interview_id": payload.interview_id}, {"$set": {"status": "completed", "qa": qa, "feedback": res.get("feedback")}})
        return {"finished": True, "feedback": res.get("feedback")}
    else:
        question = res.get("question", "Thank you, that's all for today.")
        qa.append({"q": question})
        await db.interviews.update_one({"interview_id": payload.interview_id}, {"$set": {"qa": qa}})
        return {"finished": False, "next_question": question}

@api.get("/interview/history")
async def interview_history(user: dict = Depends(get_current_user)):
    interviews = await db.interviews.find({"user_id": user["user_id"]}, {"_id": 0}).sort("created_at", -1).to_list(50)
    return {"interviews": interviews}

# ---------- voice-interview ----------
@api.get("/voice-interview/presets")
async def voice_interview_presets(user: dict = Depends(get_current_user)):
    voices = [{"key": k, **v} for k, v in VOICE_PRESETS.items()]
    personalities = [{"key": k, "label": k.replace("_", " ").title(), "desc": v} for k, v in PERSONALITIES.items()]
    types = ["technical", "hr", "recruiter", "coding", "behavioral", "system_design"]
    difficulties = ["beginner", "intermediate", "advanced", "faang"]
    return {"voices": voices, "personalities": personalities, "types": types, "difficulties": difficulties}

@api.post("/voice-interview/start")
async def voice_interview_start(payload: VoiceInterviewStart, user: dict = Depends(get_current_user)):
    interview_id = f"vint_{uuid.uuid4().hex[:12]}"
    # Get resume if needed
    resume_context = ""
    if payload.use_resume:
        resume = await db.resumes.find_one({"user_id": user["user_id"]}, {"_id": 0}, sort=[("created_at", -1)])
        if resume:
            resume_context = f"Candidate's resume: {resume.get('raw_text', '')}"
    sys_prompt = (
        f"You are a {payload.personality} {payload.interview_type} interviewer for the role of {payload.role}. "
        f"{payload.difficulty} difficulty. Ask the first question. Return STRICT JSON with 'question': string"
    )
    res = await _gemini_json(sys_prompt, f"{resume_context} Ask the first interview question.")
    question = res.get("question", "Tell me about yourself.")
    doc = {
        "interview_id": interview_id,
        "user_id": user["user_id"],
        "role": payload.role,
        "interview_type": payload.interview_type,
        "difficulty": payload.difficulty,
        "personality": payload.personality,
        "voice": payload.voice,
        "use_resume": payload.use_resume,
        "status": "active",
        "created_at": now_utc().isoformat(),
        "turns": [{"q": question}],
        "emotion_log": []
    }
    await db.voice_interviews.insert_one(dict(doc))
    return jsonable(doc)

@api.get("/voice-interview")
async def voice_interview_list(user: dict = Depends(get_current_user)):
    interviews = await db.voice_interviews.find({"user_id": user["user_id"]}, {"_id": 0}).sort("created_at", -1).to_list(50)
    return {"interviews": interviews}

@api.get("/voice-interview/{interview_id}")
async def voice_interview_get(interview_id: str, user: dict = Depends(get_current_user)):
    interview = await db.voice_interviews.find_one({"interview_id": interview_id, "user_id": user["user_id"]}, {"_id": 0})
    if not interview:
        raise HTTPException(status_code=404, detail="Interview not found")
    return interview

@api.post("/voice-interview/answer")
async def voice_interview_answer(payload: VoiceInterviewAnswer, user: dict = Depends(get_current_user)):
    interview = await db.voice_interviews.find_one({"interview_id": payload.interview_id, "user_id": user["user_id"]})
    if not interview:
        raise HTTPException(status_code=404, detail="Interview not found")
    turns = interview["turns"]
    # Update last turn
    turns[-1]["a"] = payload.transcript
    turns[-1]["confidence"] = payload.confidence
    turns[-1]["duration_s"] = payload.duration_s
    turns[-1]["wps"] = payload.wps
    turns[-1]["filler_words"] = payload.filler_words
    turns[-1]["code"] = payload.code
    turns[-1]["language"] = payload.language

    # Check if we're done (6 turns)
    if len(turns) >= 6:
        # Generate report
        sys_prompt = (
            "You are an interview evaluator. Evaluate this interview and return STRICT JSON: "
            "{"
            "overall: int 0-100, "
            "technical_score: int 0-100, "
            "communication_score: int 0-100, "
            "confidence_score: int 0-100, "
            "problem_solving_score: int 0-100, "
            "clarity_score: int 0-100, "
            "strengths: [string], "
            "improvements: [string], "
            "recommended_topics: [string], "
            "suggested_projects: [string], "
            "summary: string"
            "}"
        )
        report = await _gemini_json(sys_prompt, json.dumps(turns))
        await db.voice_interviews.update_one(
            {"interview_id": payload.interview_id},
            {"$set": {"status": "completed", "turns": turns, "report": report}}
        )
        return {"finished": True, "report": report}
    else:
        # Ask next question
        sys_prompt = (
            f"You are a {interview['personality']} {interview['interview_type']} interviewer for {interview['role']}. "
            f"Ask the next question. Return STRICT JSON with 'question': string"
        )
        res = await _gemini_json(sys_prompt, json.dumps(turns))
        question = res.get("question", "Thank you, that's all for today.")
        turns.append({"q": question})
        await db.voice_interviews.update_one(
            {"interview_id": payload.interview_id},
            {"$set": {"turns": turns}}
        )
        return {"finished": False, "next_question": question}

@api.post("/voice-interview/emotion")
async def voice_interview_emotion(payload: EmotionFrame, user: dict = Depends(get_current_user)):
    interview = await db.voice_interviews.find_one({"interview_id": payload.interview_id, "user_id": user["user_id"]})
    if not interview:
        raise HTTPException(status_code=404, detail="Interview not found")
    emotion_entry = {
        "turn_index": payload.turn_index,
        "emotions": payload.emotions,
        "ts": now_utc().isoformat()
    }
    await db.voice_interviews.update_one(
        {"interview_id": payload.interview_id},
        {"$push": {"emotion_log": emotion_entry}}
    )
    return {"ok": True}

@api.get("/voice-interview/{interview_id}/pdf")
async def voice_interview_pdf(interview_id: str, user: dict = Depends(get_current_user)):
    interview = await db.voice_interviews.find_one({"interview_id": interview_id, "user_id": user["user_id"]})
    if not interview:
        raise HTTPException(status_code=404, detail="Interview not found")
    if interview.get("status") != "completed":
        raise HTTPException(status_code=400, detail="Interview not yet completed")
    
    # Generate PDF
    buffer = io.BytesIO()
    doc_pdf = SimpleDocTemplate(buffer, pagesize=LETTER)
    styles = getSampleStyleSheet()
    story = []
    story.append(Paragraph(f"Interview Report: {interview['role']}", styles['Title']))
    story.append(Paragraph(f"Date: {interview['created_at']}", styles['Heading2']))
    report = interview.get("report", {})
    story.append(Paragraph(f"Overall Score: {report.get('overall', 0)}/100", styles['Heading3']))
    
    # Scores
    scores = [
        ("Technical Score", report.get('technical_score')),
        ("Communication Score", report.get('communication_score')),
        ("Confidence Score", report.get('confidence_score')),
        ("Problem Solving Score", report.get('problem_solving_score')),
        ("Clarity Score", report.get('clarity_score'))
    ]
    for name, score in scores:
        story.append(Paragraph(f"{name}: {score or 0}/100", styles['Normal']))
    story.append(Spacer(1, 12))
    
    # Summary
    story.append(Paragraph("Summary", styles['Heading3']))
    story.append(Paragraph(report.get('summary', ""), styles['Normal']))
    story.append(Spacer(1, 12))
    
    # Strengths
    story.append(Paragraph("Strengths", styles['Heading3']))
    for s in report.get('strengths', []):
        story.append(Paragraph(f"- {s}", styles['Normal']))
    story.append(Spacer(1, 12))
    
    # Improvements
    story.append(Paragraph("Areas for Improvement", styles['Heading3']))
    for i in report.get('improvements', []):
        story.append(Paragraph(f"- {i}", styles['Normal']))
    story.append(Spacer(1, 12))
    
    # QA
    story.append(PageBreak())
    story.append(Paragraph("Interview Questions & Answers", styles['Heading2']))
    for i, turn in enumerate(interview.get('turns', [])):
        story.append(Paragraph(f"Q{i+1}: {turn.get('q', '')}", styles['Heading3']))
        story.append(Paragraph(f"A{i+1}: {turn.get('a', '')}", styles['Normal']))
        story.append(Spacer(1, 12))
    
    doc_pdf.build(story)
    buffer.seek(0)
    return StreamingResponse(buffer, media_type="application/pdf", headers={"Content-Disposition": f"attachment; filename=interview_report_{interview_id}.pdf"})

# ---------- voice/stt ----------
@api.post("/voice/stt")
async def voice_stt(file: UploadFile = File(...), user: dict = Depends(get_current_user)):
    if not DEEPGRAM_API_KEY:
        raise HTTPException(status_code=503, detail="Deepgram not configured")
    audio_bytes = await file.read()
    try:
        async with httpx.AsyncClient() as client:
            response = await client.post(
                "https://api.deepgram.com/v1/listen",
                headers={"Authorization": f"Token {DEEPGRAM_API_KEY}", "Content-Type": "audio/*"},
                data=audio_bytes,
                timeout=60
            )
            response.raise_for_status()
            data = response.json()
            transcript = data.get("results", {}).get("channels", [{}])[0].get("alternatives", [{}])[0].get("transcript", "")
            confidence = data.get("results", {}).get("channels", [{}])[0].get("alternatives", [{}])[0].get("confidence", 0)
            return {"transcript": transcript, "confidence": int(confidence * 100)}
    except Exception as e:
        log.error(f"Deepgram error: {e}")
        raise HTTPException(status_code=502, detail="STT failed")

# ---------- voice/stt-ws (WebSocket) ----------
@app.websocket("/api/voice/stt-ws")
async def voice_stt_ws(websocket: WebSocket):
    # Note: This is a placeholder. For real implementation, you'd proxy to Deepgram WS API
    await websocket.accept()
    try:
        while True:
            data = await websocket.receive_bytes()
            # Echo back for now
            await websocket.send_json({"type": "transcript", "text": "", "is_final": False})
    except WebSocketDisconnect:
        pass
    except Exception as e:
        log.error(f"STT WS error: {e}")

# ---------- code/evaluate ----------
class CodeEvaluateIn(BaseModel):
    code: str
    language: str
    problem: str

@api.post("/code/evaluate")
async def code_evaluate(payload: CodeEvaluateIn, user: dict = Depends(get_current_user)):
    sys_prompt = (
        "You are a senior software engineer. Evaluate this code. Return STRICT JSON: "
        "{"
        "correctness: int 0-100, "
        "overall: int 0-100, "
        "complexity: string, "
        "strengths: [string], "
        "improvements: [string]"
        "}"
    )
    res = await _gemini_json(sys_prompt, f"Problem: {payload.problem}\n\nCode:\n{payload.code}")
    return res

# ---------- resume/rewrite ----------
@api.post("/resume/rewrite")
async def resume_rewrite(user: dict = Depends(get_current_user)):
    # Step 1: Get all relevant data for resume
    resume_doc = await db.resumes.find_one({"user_id": user["user_id"]}, {"_id": 0}, sort=[("created_at", -1)])
    if not resume_doc:
        raise HTTPException(status_code=400, detail="No resume uploaded")
    
    portfolio_doc = await db.portfolios.find_one({"user_id": user["user_id"]}, {"_id": 0})
    
    # Prepare context for Gemini
    raw_text = resume_doc.get('raw_text', '')
    parsed = resume_doc.get('parsed', {})
    user_profile = {
        'name': user.get('name', 'Your Name'),
        'email': user.get('email', ''),
        'github_url': user.get('github_url', ''),
        'portfolio_url': user.get('portfolio_url', ''),
        'skills': user.get('skills', []),
        'career_goals': user.get('career_goals', ''),
        'education': user.get('education', ''),
        'degree': user.get('degree', ''),
        'graduation_year': user.get('graduation_year', ''),
    }
    
    # Step 2: Use Gemini to generate structured, ATS-optimized resume content
    sys_prompt = (
        "You are an expert professional resume writer and ATS optimization specialist. "
        "Generate a complete, ATS-optimized resume. "
        "Return your response ONLY as valid JSON with these keys (NO OTHER TEXT):\n"
        "  'contact': {'name': string, 'email': string, 'phone': string, 'linkedin': string, 'github': string, 'portfolio': string},\n"
        "  'professional_summary': string,\n"
        "  'skills': {'technical': [string], 'soft': [string], 'tools': [string]},\n"
        "  'projects': [{'name': string, 'description': string, 'technologies': [string]}],\n"
        "  'education': [{'degree': string, 'institution': string, 'year': string}],\n"
        "  'certifications': [string],\n"
        "  'experience': [{'role': string, 'company': string, 'duration': string, 'achievements': [string]}]"
    )
    
    context = f"User profile: {json.dumps(user_profile)}\n\nOriginal resume text: {raw_text}\n\nParsed data: {json.dumps(parsed)}"
    if portfolio_doc:
        context += f"\n\nPortfolio data: {json.dumps(portfolio_doc)}"
        
    # Try to get JSON, if failed use parsed data from uploaded resume
    try:
        ai_content = await _gemini_json(sys_prompt, context)
    except Exception as e:
        log.warning(f"Failed to get Gemini JSON, using fallback: {e}")
        ai_content = {}
    
    # Build fallback content from parsed resume data
    if not ai_content:
        ai_content = {
            'contact': {
                'name': user.get('name', 'Your Name'),
                'email': user.get('email', ''),
                'phone': '',
                'linkedin': '',
                'github': user.get('github_url', ''),
                'portfolio': user.get('portfolio_url', '')
            },
            'professional_summary': f"Experienced professional with skills in {', '.join(user.get('skills', []))}.",
            'skills': {
                'technical': parsed.get('skills', []) or user.get('skills', []),
                'soft': ['Teamwork', 'Communication'],
                'tools': []
            },
            'projects': parsed.get('projects', []) or [],
            'education': parsed.get('education', []) or [],
            'certifications': parsed.get('certifications', []) or [],
            'experience': parsed.get('experience', []) or []
        }
    
    # Ensure all keys exist
    ai_content.setdefault('contact', {})
    ai_content.setdefault('professional_summary', '')
    ai_content.setdefault('skills', {})
    ai_content.setdefault('projects', [])
    ai_content.setdefault('education', [])
    ai_content.setdefault('certifications', [])
    ai_content.setdefault('experience', [])
    
    # Step3: Generate professional DOCX
    buffer = io.BytesIO()
    doc = docx_lib.Document()
    
    # Set margins for ATS compliance
    sections = doc.sections
    for section in sections:
        section.top_margin = docx_lib.shared.Inches(0.75)
        section.bottom_margin = docx_lib.shared.Inches(0.75)
        section.left_margin = docx_lib.shared.Inches(0.75)
        section.right_margin = docx_lib.shared.Inches(0.75)
        
    # Normalize fonts
    def set_font(paragraph, name='Calibri', size=11, bold=False):
        for run in paragraph.runs:
            run.font.name = name
            run.font.size = docx_lib.shared.Pt(size)
            run.font.bold = bold
            # For compatibility with Word on macOS
            run._element.rPr.rFonts.set(docx_lib.oxml.ns.qn('w:eastAsia'), name)
            
    # Contact Information Section
    contact = ai_content.get('contact', {})
    contact_name = contact.get('name', user.get('name', 'Your Name'))
    header_para = doc.add_paragraph()
    header_run = header_para.add_run(contact_name)
    header_run.font.name = 'Calibri'
    header_run.font.size = docx_lib.shared.Pt(16)
    header_run.font.bold = True
    header_para.alignment = docx_lib.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    header_run._element.rPr.rFonts.set(docx_lib.oxml.ns.qn('w:eastAsia'), 'Calibri')
    
    contact_line = []
    if contact.get('email'): contact_line.append(contact['email'])
    if contact.get('phone'): contact_line.append(contact['phone'])
    if contact.get('linkedin'): contact_line.append(contact['linkedin'])
    if contact.get('github'): contact_line.append(contact['github'])
    if contact.get('portfolio'): contact_line.append(contact['portfolio'])
    
    if contact_line:
        contact_para = doc.add_paragraph()
        contact_para.add_run(' | '.join(contact_line))
        contact_para.alignment = docx_lib.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        set_font(contact_para, size=11)
        
    doc.add_paragraph()  # Add spacing
    
    # Professional Summary
    if ai_content.get('professional_summary'):
        summary_heading = doc.add_paragraph()
        summary_heading.add_run('PROFESSIONAL SUMMARY').bold = True
        set_font(summary_heading, size=12)
        summary_para = doc.add_paragraph(ai_content['professional_summary'])
        set_font(summary_para, size=11)
        summary_para.paragraph_format.line_spacing = 1.15
        doc.add_paragraph()
        
    # Skills Section
    skills = ai_content.get('skills', {})
    if skills:
        skills_heading = doc.add_paragraph()
        skills_heading.add_run('SKILLS').bold = True
        set_font(skills_heading, size=12)
        
        skill_items = []
        if skills.get('technical'):
            skill_items.append(f"Technical: {', '.join(skills['technical'])}")
        if skills.get('tools'):
            skill_items.append(f"Tools: {', '.join(skills['tools'])}")
        if skills.get('soft'):
            skill_items.append(f"Soft: {', '.join(skills['soft'])}")
            
        skills_para = doc.add_paragraph()
        skills_para.add_run('; '.join(skill_items))
        set_font(skills_para, size=11)
        doc.add_paragraph()
        
    # Experience Section
    experience = ai_content.get('experience', [])
    if experience:
        exp_heading = doc.add_paragraph()
        exp_heading.add_run('PROFESSIONAL EXPERIENCE').bold = True
        set_font(exp_heading, size=12)
        
        for exp in experience:
            role = exp.get('role', '')
            company = exp.get('company', '')
            duration = exp.get('duration', '')
            
            role_para = doc.add_paragraph()
            role_run = role_para.add_run(f"{role}")
            role_run.bold = True
            set_font(role_para, size=11)
            role_para.paragraph_format.space_after = docx_lib.shared.Pt(0)
            
            company_para = doc.add_paragraph()
            company_para.add_run(f"{company} | {duration}").italic = True
            set_font(company_para, size=11)
            company_para.paragraph_format.space_after = docx_lib.shared.Pt(0)
            
            for ach in exp.get('achievements', []):
                ach_para = doc.add_paragraph(ach, style='List Bullet')
                set_font(ach_para, size=11)
                ach_para.paragraph_format.space_after = docx_lib.shared.Pt(0)
                
            doc.add_paragraph()
            
    # Projects Section
    projects = ai_content.get('projects', [])
    if projects:
        proj_heading = doc.add_paragraph()
        proj_heading.add_run('PROJECTS').bold = True
        set_font(proj_heading, size=12)
        
        for proj in projects:
            name = proj.get('name', '')
            techs = ', '.join(proj.get('technologies', []))
            desc = proj.get('description', '')
            
            proj_name_para = doc.add_paragraph()
            proj_name_para.add_run(name).bold = True
            if techs:
                proj_name_para.add_run(f" — {techs}").italic = True
            set_font(proj_name_para, size=11)
            proj_name_para.paragraph_format.space_after = docx_lib.shared.Pt(0)
            
            desc_para = doc.add_paragraph(desc)
            set_font(desc_para, size=11)
            doc.add_paragraph()
            
    # Education Section
    education = ai_content.get('education', [])
    if education:
        edu_heading = doc.add_paragraph()
        edu_heading.add_run('EDUCATION').bold = True
        set_font(edu_heading, size=12)
        
        for edu in education:
            degree = edu.get('degree', '')
            institution = edu.get('institution', '')
            year = edu.get('year', '')
            
            edu_para = doc.add_paragraph()
            edu_para.add_run(f"{degree}").bold = True
            set_font(edu_para, size=11)
            edu_para.paragraph_format.space_after = docx_lib.shared.Pt(0)
            
            inst_para = doc.add_paragraph(f"{institution} | {year}")
            set_font(inst_para, size=11)
            doc.add_paragraph()
            
    # Certifications Section
    certifications = ai_content.get('certifications', [])
    if certifications:
        cert_heading = doc.add_paragraph()
        cert_heading.add_run('CERTIFICATIONS').bold = True
        set_font(cert_heading, size=12)
        
        for cert in certifications:
            cert_para = doc.add_paragraph(cert, style='List Bullet')
            set_font(cert_para, size=11)
        doc.add_paragraph()
        
    # Save document
    doc.save(buffer)
    buffer.seek(0)
    
    # Return file for download
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename=ATS_Optimized_Resume.docx"}
    )

# ---------- career-twin ----------
@api.post("/career-twin/brief")
async def career_twin_brief(user: dict = Depends(get_current_user)):
    week_of = now_utc().strftime("%Y-%m-%d")
    sys_prompt = (
        "You are an AI career twin. Generate a weekly brief. Return STRICT JSON: "
        "{"
        "greeting: string, "
        "this_week_focus: string, "
        "market_signals: [string], "
        "recommended_action: string, "
        "opportunities: [{title: string, why_now: string}]"
        "}"
    )
    profile = json.dumps(user)
    brief = await _gemini_json(sys_prompt, profile)
    doc = {
        "user_id": user["user_id"],
        "week_of": week_of,
        "brief": brief,
        "created_at": now_utc().isoformat()
    }
    await db.career_twin.insert_one(dict(doc))
    return jsonable(doc)

@api.get("/career-twin")
async def career_twin_get(user: dict = Depends(get_current_user)):
    latest = await db.career_twin.find_one({"user_id": user["user_id"]}, {"_id": 0}, sort=[("created_at", -1)])
    return latest or {}

# ---------- public-profile ----------
@api.get("/public-profile/me")
async def public_profile_me(user: dict = Depends(get_current_user)):
    profile = await db.public_profiles.find_one({"user_id": user["user_id"]}, {"_id": 0})
    return profile or {}

@api.post("/public-profile/publish")
async def public_profile_publish(payload: PublishIn, user: dict = Depends(get_current_user)):
    # Check if slug exists
    existing = await db.public_profiles.find_one({"slug": payload.slug})
    if existing and existing["user_id"] != user["user_id"]:
        raise HTTPException(status_code=400, detail="Slug already taken")
    # Get user data
    career_rec = await db.career_recommendations.find_one({"user_id": user["user_id"]}, {"_id": 0}, sort=[("created_at", -1)])
    portfolio = await db.portfolios.find_one({"user_id": user["user_id"]}, {"_id": 0})
    doc = {
        "user_id": user["user_id"],
        "slug": payload.slug,
        "name": user.get("name", ""),
        "picture": user.get("picture", ""),
        "headline": payload.headline or "",
        "bio": payload.bio or "",
        "skills": user.get("skills", []),
        "show_resume": payload.show_resume,
        "show_portfolio": payload.show_portfolio,
        "top_careers": career_rec.get("careers", [])[:5] if career_rec else [],
        "portfolio": portfolio if portfolio and payload.show_portfolio else None,
        "created_at": now_utc().isoformat(),
        "updated_at": now_utc().isoformat()
    }
    await db.public_profiles.update_one({"user_id": user["user_id"]}, {"$set": doc}, upsert=True)
    return jsonable(doc)

# ---------- public profile (unauthenticated) ----------
@api.get("/public/{slug}")
async def public_profile(slug: str):
    profile = await db.public_profiles.find_one({"slug": slug}, {"_id": 0, "user_id": 0})
    if not profile:
        raise HTTPException(status_code=404, detail="Profile not found")
    return profile

# ---------- startup / shutdown / cors ----------
app.include_router(api)
app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=os.environ.get("CORS_ORIGINS", "*").split(","),
    allow_methods=["*"],
    allow_headers=["*"],
)


@app.on_event("startup")
async def startup():
    # Indexes for performance — all data lives in MongoDB Atlas
    try:
        await db.users.create_index("user_id", unique=True)
        await db.users.create_index("email", unique=True)
        await db.user_sessions.create_index("session_token", unique=True)
        await db.user_sessions.create_index("user_id")
        await db.resumes.create_index([("user_id", 1), ("created_at", -1)])
        await db.career_recommendations.create_index("user_id")
        await db.skill_gaps.create_index([("user_id", 1), ("target_role", 1)])
        await db.roadmaps.create_index([("user_id", 1), ("created_at", -1)])
        await db.voice_interviews.create_index([("user_id", 1), ("created_at", -1)])
        await db.interviews.create_index([("user_id", 1), ("created_at", -1)])
        await db.chat_messages.create_index([("user_id", 1), ("ts", 1)])
        await db.chat_memory.create_index("user_id")
        await db.portfolios.create_index("user_id", unique=True)
        await db.career_twin.create_index([("user_id", 1), ("week_of", -1)])
        await db.public_profiles.create_index("slug", unique=True)
        log.info("Mongo indexes ensured")
    except Exception as e:
        log.warning("index init: %s", e)


@app.on_event("shutdown")
async def shutdown():
    client.close()

