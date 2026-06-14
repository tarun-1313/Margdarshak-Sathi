"""
Premium ATS Resume Generator Service
Generates professional, ATS-optimized resumes using python-docx
"""
import os
import io
import json
import re
import traceback
from datetime import datetime
from typing import Dict, List, Optional, Any
from dataclasses import dataclass, asdict
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


@dataclass
class PersonalInfo:
    """Personal information data class"""
    name: str
    title: str
    email: str
    phone: str
    location: str
    linkedin: str = ""
    github: str = ""
    website: str = ""
    summary: str = ""


@dataclass
class Experience:
    """Work experience data class"""
    role: str
    company: str
    location: str
    duration: str
    achievements: List[str]
    technologies: List[str]


@dataclass
class Education:
    """Education data class"""
    degree: str
    school: str
    year: str
    gpa: str = ""


@dataclass
class Project:
    """Project data class"""
    name: str
    description: str
    technologies: List[str]
    github: str = ""
    demo: str = ""


@dataclass
class Certification:
    """Certification data class"""
    name: str
    issuer: str
    year: str


@dataclass
class ResumeData:
    """Complete resume data"""
    personal: PersonalInfo
    skills: Dict[str, List[Dict[str, Any]]]
    experience: List[Experience]
    education: List[Education]
    projects: List[Project]
    certifications: List[Certification]


class ResumeGenerator:
    """Premium resume generator using python-docx"""
    
    # Professional color schemes
    COLORS = {
        'modern_ai': {
            'primary': RGBColor(139, 92, 246),      # Purple
            'secondary': RGBColor(59, 130, 246),    # Blue
            'accent': RGBColor(163, 255, 18),       # Green
            'text': RGBColor(31, 41, 55),           # Dark gray
            'light': RGBColor(107, 114, 128),       # Gray
        },
        'faang': {
            'primary': RGBColor(37, 99, 235),       # Blue
            'secondary': RGBColor(29, 78, 216),     # Dark blue
            'accent': RGBColor(16, 185, 129),       # Green
            'text': RGBColor(17, 24, 39),           # Almost black
            'light': RGBColor(75, 85, 99),            # Gray
        },
        'startup': {
            'primary': RGBColor(245, 158, 11),      # Orange
            'secondary': RGBColor(239, 68, 68),     # Red
            'accent': RGBColor(16, 185, 129),       # Green
            'text': RGBColor(31, 41, 55),           # Dark gray
            'light': RGBColor(107, 114, 128),       # Gray
        },
    }
    
    def __init__(self, template: str = "modern_ai"):
        self.template = template
        self.colors = self.COLORS.get(template, self.COLORS['modern_ai'])
        self.doc = Document()
        self.setup_document()
    
    def setup_document(self):
        """Setup document margins and styles"""
        # Set margins (in inches)
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.6)
            section.right_margin = Inches(0.6)
        
        # Setup styles
        self.setup_styles()
    
    def setup_styles(self):
        """Setup custom paragraph styles"""
        styles = self.doc.styles
        
        # Name style
        name_style = styles.add_style('ResumeName', WD_STYLE_TYPE.PARAGRAPH)
        name_style.font.name = 'Calibri'
        name_style.font.size = Pt(26)
        name_style.font.bold = True
        name_style.font.color.rgb = self.colors['text']
        name_style.paragraph_format.space_after = Pt(2)
        
        # Title style
        title_style = styles.add_style('ResumeTitle', WD_STYLE_TYPE.PARAGRAPH)
        title_style.font.name = 'Calibri'
        title_style.font.size = Pt(12)
        title_style.font.color.rgb = self.colors['primary']
        title_style.paragraph_format.space_after = Pt(6)
        
        # Section header style
        section_style = styles.add_style('ResumeSection', WD_STYLE_TYPE.PARAGRAPH)
        section_style.font.name = 'Calibri'
        section_style.font.size = Pt(13)
        section_style.font.bold = True
        section_style.font.color.rgb = self.colors['text']
        section_style.paragraph_format.space_before = Pt(10)
        section_style.paragraph_format.space_after = Pt(4)
        
        # Body text style
        body_style = styles.add_style('ResumeBody', WD_STYLE_TYPE.PARAGRAPH)
        body_style.font.name = 'Calibri'
        body_style.font.size = Pt(10)
        body_style.font.color.rgb = self.colors['text']
        body_style.paragraph_format.space_after = Pt(2)
        body_style.paragraph_format.line_spacing = 1.15
    
    def add_header(self, personal: PersonalInfo):
        """Add resume header with name and contact info"""
        # Name
        name_para = self.doc.add_paragraph(personal.name, style='ResumeName')
        name_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Title
        title_para = self.doc.add_paragraph(personal.title, style='ResumeTitle')
        title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Contact info in one line
        contact_parts = []
        if personal.email:
            contact_parts.append(personal.email)
        if personal.phone:
            contact_parts.append(personal.phone)
        if personal.location:
            contact_parts.append(personal.location)
        if personal.linkedin:
            contact_parts.append(personal.linkedin)
        
        if contact_parts:
            contact_para = self.doc.add_paragraph(' • '.join(contact_parts))
            contact_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            contact_para.runs[0].font.size = Pt(9)
            contact_para.runs[0].font.color.rgb = self.colors['light']
            contact_para.paragraph_format.space_after = Pt(8)
        
        # Summary
        if personal.summary:
            summary_para = self.doc.add_paragraph(personal.summary, style='ResumeBody')
            summary_para.paragraph_format.space_after = Pt(8)
    
    def add_section_header(self, title: str):
        """Add a section header with underline"""
        # Add section title
        section_para = self.doc.add_paragraph(title.upper(), style='ResumeSection')
        
        # Add underline
        underline_para = self.doc.add_paragraph()
        underline_para.paragraph_format.space_after = Pt(6)
        run = underline_para.add_run('_' * 60)
        run.font.size = Pt(8)
        run.font.color.rgb = self.colors['primary']
    
    def add_skills(self, skills: Dict[str, List[Dict[str, Any]]]):
        """Add skills section"""
        if not skills:
            return
        
        self.add_section_header("Technical Skills")
        
        for category, items in skills.items():
            skill_para = self.doc.add_paragraph(style='ResumeBody')
            
            # Category name in bold
            category_run = skill_para.add_run(f"{category}: ")
            category_run.bold = True
            category_run.font.size = Pt(10)
            
            # Skills list
            skills_list = [item['name'] if isinstance(item, dict) else item for item in items]
            skills_run = skill_para.add_run(', '.join(skills_list))
            skills_run.font.size = Pt(10)
            
            skill_para.paragraph_format.space_after = Pt(3)
    
    def add_experience(self, experiences: List[Experience]):
        """Add work experience section"""
        if not experiences:
            return
        
        self.add_section_header("Professional Experience")
        
        for exp in experiences:
            # Role and Company
            header_para = self.doc.add_paragraph()
            header_para.paragraph_format.space_after = Pt(2)
            
            role_run = header_para.add_run(exp.role)
            role_run.bold = True
            role_run.font.size = Pt(11)
            role_run.font.color.rgb = self.colors['text']
            
            company_run = header_para.add_run(f" | {exp.company}")
            company_run.font.size = Pt(10)
            company_run.font.color.rgb = self.colors['light']
            
            # Duration and Location
            details_para = self.doc.add_paragraph()
            details_para.paragraph_format.space_after = Pt(4)
            details_run = details_para.add_run(f"{exp.duration} | {exp.location}")
            details_run.italic = True
            details_run.font.size = Pt(9)
            details_run.font.color.rgb = self.colors['light']
            
            # Achievements
            for achievement in exp.achievements:
                bullet_para = self.doc.add_paragraph(style='List Bullet')
                bullet_para.paragraph_format.left_indent = Inches(0.2)
                bullet_para.paragraph_format.space_after = Pt(2)
                bullet_run = bullet_para.add_run(achievement)
                bullet_run.font.size = Pt(9)
            
            # Technologies
            if exp.technologies:
                tech_para = self.doc.add_paragraph()
                tech_para.paragraph_format.space_after = Pt(6)
                tech_run = tech_para.add_run(f"Technologies: {', '.join(exp.technologies)}")
                tech_run.italic = True
                tech_run.font.size = Pt(8)
                tech_run.font.color.rgb = self.colors['light']
    
    def add_projects(self, projects: List[Project]):
        """Add projects section"""
        if not projects:
            return
        
        self.add_section_header("Featured Projects")
        
        for project in projects:
            # Project Name
            name_para = self.doc.add_paragraph()
            name_para.paragraph_format.space_after = Pt(2)
            name_run = name_para.add_run(project.name)
            name_run.bold = True
            name_run.font.size = Pt(10)
            
            # Description
            desc_para = self.doc.add_paragraph(project.description, style='ResumeBody')
            desc_para.paragraph_format.space_after = Pt(2)
            
            # Technologies
            tech_para = self.doc.add_paragraph()
            tech_para.paragraph_format.space_after = Pt(4)
            tech_run = tech_para.add_run(f"Tech Stack: {', '.join(project.technologies)}")
            tech_run.italic = True
            tech_run.font.size = Pt(8)
            tech_run.font.color.rgb = self.colors['light']
    
    def add_education(self, educations: List[Education]):
        """Add education section"""
        if not educations:
            return
        
        self.add_section_header("Education")
        
        for edu in educations:
            # Degree and School
            header_para = self.doc.add_paragraph()
            header_para.paragraph_format.space_after = Pt(2)
            
            degree_run = header_para.add_run(edu.degree)
            degree_run.bold = True
            degree_run.font.size = Pt(10)
            
            school_run = header_para.add_run(f" | {edu.school}")
            school_run.font.size = Pt(9)
            school_run.font.color.rgb = self.colors['light']
            
            # Year and GPA
            details_para = self.doc.add_paragraph()
            details_para.paragraph_format.space_after = Pt(4)
            details_text = edu.year
            if edu.gpa:
                details_text += f" | GPA: {edu.gpa}"
            details_run = details_para.add_run(details_text)
            details_run.italic = True
            details_run.font.size = Pt(8)
            details_run.font.color.rgb = self.colors['light']
    
    def add_certifications(self, certifications: List[Certification]):
        """Add certifications section"""
        if not certifications:
            return
        
        self.add_section_header("Certifications")
        
        for cert in certifications:
            cert_para = self.doc.add_paragraph(style='ResumeBody')
            cert_para.paragraph_format.space_after = Pt(2)
            
            name_run = cert_para.add_run(cert.name)
            name_run.bold = True
            
            details_run = cert_para.add_run(f" | {cert.issuer} | {cert.year}")
            details_run.font.size = Pt(9)
            details_run.font.color.rgb = self.colors['light']
    
    def generate_resume(self, data: ResumeData) -> Document:
        """Generate complete resume document"""
        # Add header with personal info
        self.add_header(data.personal)
        
        # Add skills section
        self.add_skills(data.skills)
        
        # Add experience section
        self.add_experience(data.experience)
        
        # Add projects section
        self.add_projects(data.projects)
        
        # Add education section
        self.add_education(data.education)
        
        # Add certifications if any
        if data.certifications:
            self.add_certifications(data.certifications)
        
        return self.doc
    
    def save_to_bytes(self) -> bytes:
        """Save document to bytes buffer"""
        buffer = io.BytesIO()
        self.doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()


def generate_premium_resume(resume_data: Dict[str, Any], template: str = "modern_ai") -> bytes:
    """
    Generate a premium ATS-optimized resume
    
    Args:
        resume_data: Dictionary containing all resume information
        template: Template style ("modern_ai", "faang", "startup", "research")
    
    Returns:
        bytes: DOCX file content
    """
    try:
        # Convert dict to dataclasses
        personal = PersonalInfo(**resume_data['personal'])
        
        experiences = [Experience(**exp) for exp in resume_data.get('experience', [])]
        educations = [Education(**edu) for edu in resume_data.get('education', [])]
        projects = [Project(**proj) for proj in resume_data.get('projects', [])]
        certifications = [Certification(**cert) for cert in resume_data.get('certifications', [])]
        
        data = ResumeData(
            personal=personal,
            skills=resume_data.get('skills', {}),
            experience=experiences,
            education=educations,
            projects=projects,
            certifications=certifications
        )
        
        # Generate resume
        generator = ResumeGenerator(template=template)
        generator.generate_resume(data)
        
        # Return bytes
        return generator.save_to_bytes()
        
    except Exception as e:
        print(f"Error generating resume: {str(e)}")
        traceback.print_exc()
        raise


# Example usage
if __name__ == "__main__":
    # Example resume data
    example_data = {
        "personal": {
            "name": "Alex Johnson",
            "title": "AI Engineer | Generative AI | LLM Applications",
            "email": "alex@example.com",
            "phone": "+1 (555) 123-4567",
            "location": "San Francisco, CA",
            "linkedin": "linkedin.com/in/alexjohnson",
            "github": "github.com/alexjohnson",
            "website": "alexjohnson.dev",
            "summary": "Building intelligent AI systems, multi-agent workflows, and production-grade AI applications using modern AI infrastructure."
        },
        "skills": {
            "AI & Machine Learning": [
                {"name": "TensorFlow", "level": 90},
                {"name": "PyTorch", "level": 85},
                {"name": "LangChain", "level": 95},
            ],
            "Backend": [
                {"name": "FastAPI", "level": 95},
                {"name": "Node.js", "level": 80},
            ],
        },
        "experience": [
            {
                "role": "Senior AI Engineer",
                "company": "TechCorp AI",
                "location": "San Francisco, CA",
                "duration": "2023 - Present",
                "achievements": [
                    "Built scalable AI pipelines improving processing efficiency by 35%",
                    "Led development of multi-agent workflow system using LangGraph",
                ],
                "technologies": ["Python", "LangChain", "FastAPI", "OpenAI"]
            }
        ],
        "education": [
            {
                "degree": "B.S. Computer Science",
                "school": "Stanford University",
                "year": "2019 - 2023",
                "gpa": "3.8/4.0"
            }
        ],
        "projects": [
            {
                "name": "CareerPilot AI",
                "description": "AI-powered career operating system with interviews, skill gap analysis, roadmaps, and ATS optimization.",
                "technologies": ["FastAPI", "Gemini", "MongoDB", "LangGraph"],
                "github": "#",
                "demo": "#"
            }
        ],
        "certifications": [
            {
                "name": "Deep Learning Specialization",
                "issuer": "Coursera",
                "year": "2024"
            }
        ]
    }
    
    # Generate resume
    try:
        resume_bytes = generate_premium_resume(example_data, template="modern_ai")
        
        # Save to file
        with open("generated_resume.docx", "wb") as f:
            f.write(resume_bytes)
        
        print("✅ Resume generated successfully: generated_resume.docx")
    except Exception as e:
        print(f"❌ Error: {str(e)}")
